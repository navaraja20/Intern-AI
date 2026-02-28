"""
Resume Parser – extracts clean text from PDF, DOCX, TXT.
Returns both raw text and structured sections.
"""

import io
import re
from pathlib import Path


# ── Main Entry ────────────────────────────────────────────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to correct parser. Returns clean plain text."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_bytes)
    elif ext == ".txt":
        return _clean(file_bytes.decode("utf-8", errors="ignore"))
    else:
        raise ValueError(f"Unsupported format: {ext}. Upload PDF, DOCX, or TXT.")


def validate(text: str) -> tuple[bool, str]:
    """Sanity-check that the text looks like a resume."""
    if len(text) < 200:
        return False, "Document too short – is this a resume?"
    if len(text) > 60_000:
        return False, "Document is too long. Please upload only your resume."
    indicators = ["experience","education","skill","university","project",
                  "work","degree","bachelor","master","intern","python",
                  "data","engineer","developer","analyst"]
    found = sum(1 for kw in indicators if kw in text.lower())
    if found < 2:
        return False, "This doesn't look like a resume. Please check the file."
    return True, ""


# ── Section Parser ────────────────────────────────────────────────────────────

def parse_sections(text: str) -> dict:
    """
    Extract structured sections from resume text.
    Returns dict: {summary, education, experience, skills, projects, certifications}
    """
    sections: dict[str, str] = {
        "summary": "", "education": "", "experience": "",
        "skills": "", "projects": "", "certifications": "", "other": "",
    }

    # Patterns that match common section headers
    header_patterns = {
        "summary":        r"(professional\s+summary|summary|objective|profile|about\s+me)",
        "education":      r"(education|academic|qualification)",
        "experience":     r"(experience|work\s+history|employment|professional\s+experience)",
        "skills":         r"(skill|technical\s+skill|core\s+competenc|technologies)",
        "projects":       r"(project|personal\s+project|academic\s+project|key\s+project)",
        "certifications": r"(certif|award|honor|achievement|course)",
    }

    # Split text into lines
    lines = text.split("\n")
    current_section = "other"
    section_content: dict[str, list[str]] = {k: [] for k in sections}

    for line in lines:
        stripped = line.strip()
        if not stripped:
            section_content[current_section].append("")
            continue

        # Check if this line is a section header
        matched_section = None
        for sec, pattern in header_patterns.items():
            if re.match(rf"^\s*({pattern})\s*$", stripped, re.IGNORECASE):
                matched_section = sec
                break
            # Also catch ALL CAPS headers
            if stripped.isupper() and re.search(pattern, stripped, re.IGNORECASE):
                matched_section = sec
                break

        if matched_section:
            current_section = matched_section
        else:
            section_content[current_section].append(stripped)

    for sec in sections:
        sections[sec] = "\n".join(section_content[sec]).strip()

    return sections


# ── Format-specific Parsers ───────────────────────────────────────────────────

def _parse_pdf(file_bytes: bytes) -> str:
    # Try pdfplumber first (best quality)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                try:
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                except Exception:
                    page_text = None
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return _clean("\n".join(text_parts))
    except Exception:
        pass

    # Fallback: PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        if text.strip():
            return _clean(text)
    except Exception:
        pass

    raise ValueError("Could not extract text from this PDF. "
                     "Try saving it as a plain-text PDF or uploading a DOCX/TXT instead.")


def _parse_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return _clean("\n".join(parts))
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")
    except Exception as exc:
        raise ValueError(f"Could not parse DOCX file: {exc}") from exc


def _clean(text: str) -> str:
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    text = re.sub(r" {2,}", " ", text)
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
