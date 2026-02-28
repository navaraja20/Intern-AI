"""
PDF & DOCX Generator – produces ATS-clean, professional documents.
No columns, no tables, no graphics – pure text layout.
"""

import io
import re
from config import settings


# ── PDF Generation (ReportLab) ────────────────────────────────────────────────

def _deduplicate_sections(text: str) -> str:
    """
    Remove duplicate ALL-CAPS section headers from LLM output.
    Keeps the first occurrence of each section; strips all later ones.
    Also strips common LLM meta-commentary lines.
    """
    # Known section headers
    section_keywords = {
        "PROFESSIONAL SUMMARY", "SUMMARY", "OBJECTIVE",
        "SKILLS", "TECHNICAL SKILLS", "CORE SKILLS",
        "EXPERIENCE", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE",
        "PROJECTS", "PROJECT",
        "EDUCATION", "ACADEMIC",
        "CERTIFICATIONS", "CERTIFICATIONS & AWARDS",
        "CONTACT", "CONTACT INFORMATION",
    }
    # LLM meta lines to strip entirely
    meta_patterns = [
        r"^here('s| is)\b",
        r"^below is\b",
        r"^the tailored resume\b",
        r"^resume:",
        r"^---+$",
        r"^\*\*",
        r"step \d+",
        r"final check",
        r"word count",
    ]

    seen_sections: set[str] = set()
    out_lines: list[str] = []
    skip_until_next_section = False

    for line in text.split("\n"):
        stripped = line.strip()
        upper = stripped.upper()

        # Drop meta-commentary
        if any(re.search(p, stripped, re.IGNORECASE) for p in meta_patterns):
            continue

        # Detect ALL-CAPS section header
        is_section_header = (
            stripped.isupper()
            and 3 < len(stripped) < 80
            and not re.search(r"[0-9@|]", stripped)
        )

        if is_section_header:
            # Normalise: find canonical key
            matched_key = None
            for kw in section_keywords:
                if kw in upper:
                    matched_key = kw
                    break
            matched_key = matched_key or upper

            if matched_key in seen_sections:
                # Duplicate – skip this header and all content until next section
                skip_until_next_section = True
                continue
            else:
                seen_sections.add(matched_key)
                skip_until_next_section = False
                out_lines.append(line)
                continue

        if skip_until_next_section:
            continue

        out_lines.append(line)

    return "\n".join(out_lines)


def generate_pdf(text: str, title: str = "Resume") -> bytes:
    """
    Render plain text to ATS-friendly single-page PDF.
    Auto-scales font size to fit content onto one page.
    """
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        # Clean LLM artefacts before rendering
        text = _deduplicate_sections(text)

        margin = 36  # 0.5 inch — tight but readable

        def _build_story(body_size: float):
            """Build the ReportLab story for a given body font size."""
            name_style = ParagraphStyle(
                "Name", fontSize=body_size + 4, fontName="Helvetica-Bold",
                textColor=HexColor("#1a1a2e"), alignment=TA_CENTER,
                spaceBefore=0, spaceAfter=2,
            )
            contact_style = ParagraphStyle(
                "Contact", fontSize=body_size - 1, fontName="Helvetica",
                textColor=HexColor("#444444"), alignment=TA_CENTER,
                spaceBefore=0, spaceAfter=4,
            )
            heading_style = ParagraphStyle(
                "Heading", fontSize=body_size + 0.5, fontName="Helvetica-Bold",
                textColor=HexColor("#1a1a2e"),
                spaceBefore=6, spaceAfter=1, leftIndent=0,
            )
            role_style = ParagraphStyle(
                "Role", fontSize=body_size, fontName="Helvetica-Bold",
                textColor=HexColor("#222222"),
                spaceBefore=3, spaceAfter=1, leftIndent=0,
            )
            body_style = ParagraphStyle(
                "Body", fontSize=body_size, fontName="Helvetica",
                leading=body_size * 1.35, spaceAfter=1, leftIndent=0,
            )
            bullet_style = ParagraphStyle(
                "Bullet", fontSize=body_size, fontName="Helvetica",
                leading=body_size * 1.35, spaceAfter=1,
                leftIndent=12, firstLineIndent=-8,
            )

            story = []
            lines = text.split("\n")
            first_nb = next((i for i, l in enumerate(lines) if l.strip()), 0)
            contact_zone = first_nb + 5  # first N lines may be contact info

            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()
                i += 1

                if not stripped:
                    story.append(Spacer(1, 2))
                    continue

                # Candidate name (first non-blank, not all-caps)
                if i - 1 == first_nb and not stripped.isupper():
                    story.append(Paragraph(_esc(stripped), name_style))
                    continue

                # Contact info zone
                if i - 1 <= contact_zone and re.search(r"[@|+]|linkedin|github", stripped, re.I):
                    story.append(Paragraph(_esc(stripped), contact_style))
                    continue

                # ALL CAPS section heading
                if stripped.isupper() and 3 < len(stripped) < 80 and not re.search(r"[0-9]", stripped):
                    story.append(HRFlowable(
                        width="100%", thickness=0.6,
                        color=HexColor("#4f46e5"), spaceBefore=4, spaceAfter=1,
                    ))
                    story.append(Paragraph(_esc(stripped), heading_style))
                    continue

                # Bullet point (•, -, *, —)
                if re.match(r"^[•\-\*—◦▪]\s+", stripped):
                    bullet_text = re.sub(r"^[•\-\*—◦▪]\s+", "", stripped)
                    story.append(Paragraph(f"\u2022\u00a0{_esc(bullet_text)}", bullet_style))
                    continue

                # Role/company line (contains | separator — common in job entries)
                if " | " in stripped and stripped.count("|") >= 1:
                    story.append(Paragraph(_esc(stripped), role_style))
                    continue

                story.append(Paragraph(_esc(stripped), body_style))

            return story

        # ── Auto-scale to fit 1 page ──────────────────────────────────────────
        page_w, page_h = LETTER
        usable_h = page_h - 2 * margin

        for font_size in [9.5, 9.0, 8.5, 8.0, 7.5]:
            buf = io.BytesIO()
            doc = SimpleDocTemplate(
                buf, pagesize=LETTER,
                leftMargin=margin, rightMargin=margin,
                topMargin=margin, bottomMargin=margin,
                title=title,
            )
            story = _build_story(font_size)
            doc.build(story)
            pdf_bytes = buf.getvalue()

            # Check page count by counting '%%Page' markers or doc.page
            # ReportLab sets doc.page after build
            if getattr(doc, 'page', 2) <= 1:
                return pdf_bytes  # fits on 1 page

        # Last resort: return smallest font version
        return pdf_bytes

    except ImportError:
        raise ImportError("Install reportlab: pip install reportlab")


# ── DOCX Generation (python-docx) ─────────────────────────────────────────────

def generate_docx(text: str, title: str = "Resume") -> bytes:
    """
    Render plain text to ATS-friendly DOCX.
    Returns bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Narrow margins
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.0)
            section.right_margin  = Cm(2.0)

        # Default style
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        lines = text.split("\n")
        first_nonblank_idx = next((i for i, l in enumerate(lines) if l.strip()), 0)

        for idx, line in enumerate(lines):
            stripped = line.strip()

            if not stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                continue

            # Name line
            if idx == first_nonblank_idx and not stripped.isupper():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                p.paragraph_format.space_after = Pt(3)
                continue

            # Contact line
            if re.search(r"@|(\+\d|\d{3}[.\-]\d{3})", stripped) and idx < first_nonblank_idx + 4:
                p = doc.add_paragraph(stripped)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                continue

            # Section heading
            if stripped.isupper() and 3 < len(stripped) < 70:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                # Bottom border (underline effect)
                _add_bottom_border(p)
                continue

            # Bullet
            if stripped.startswith(("•", "–", "-", "*", "▪", "◦")):
                bullet_text = re.sub(r"^[•\-\–\*▪◦]\s*", "", stripped)
                p = doc.add_paragraph(style="List Bullet")
                p.text = bullet_text
                p.paragraph_format.space_after = Pt(1)
                continue

            # Body
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(2)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")


# ── Helpers ───────────────────────────────────────────────────────────────────

def _esc(text: str) -> str:
    """Escape ReportLab XML special chars."""
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def _add_bottom_border(paragraph) -> None:
    """Add a bottom border to a DOCX paragraph (section divider)."""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "4f46e5")
        pBdr.append(bottom)
        pPr.append(pBdr)
    except Exception:
        pass  # non-critical styling
